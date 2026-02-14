import streamlit as st
from docx import Document
import random
from io import BytesIO

# --- إعدادات الصفحة ---
st.set_page_config(page_title="منصة الامتحانات - جامعة المستقبل", layout="centered")

st.image("https://upload.wikimedia.org/wikipedia/commons/thumb/e/e3/University_of_Babylon_logo.svg/1200px-University_of_Babylon_logo.svg.png", width=100) # يمكنك تغيير رابط الشعار
st.title("نظام توليد الأسئلة الامتحانية")
st.write("قم برفع ملف بنك الأسئلة (Word)، وسيقوم النظام بوضعه في القالب الرسمي للجامعة.")

# --- تحميل القالب الرسمي ---
# ملاحظة: يجب أن يكون ملف القالب موجوداً بجانب الكود في نفس المجلد
TEMPLATE_FILE = 'نموذج الاسئلة 30سؤال.docx' 

def read_questions(file):
    doc = Document(file)
    mcq_list = []
    tf_list = []
    current_mode = None
    
    lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    
    i = 0
    while i < len(lines):
        line = lines[i]
        if "# اختياري" in line:
            current_mode = "MCQ"
            i += 1
            continue
        elif "# صح وخطأ" in line:
            current_mode = "TF"
            i += 1
            continue
            
        if current_mode == "MCQ":
            if i + 5 < len(lines):
                q = lines[i]
                opts = lines[i+1:i+6]
                if not any("#" in opt for opt in opts):
                    mcq_list.append({"q": q, "opts": opts})
                    i += 6
                    continue
        elif current_mode == "TF":
            tf_list.append(line)
            i += 1
            continue
        i += 1
    return mcq_list, tf_list

def generate_exam(mcq_data, tf_data, template_path):
    doc = Document(template_path)
    random.shuffle(mcq_data)
    random.shuffle(tf_data)
    
    mcq_idx = 0
    tf_idx = 0
    
    for table in doc.tables:
        try:
            row_txt = table.rows[1].cells[0].text + table.rows[0].cells[0].text
        except: row_txt = ""

        # MCQ Logic
        if "A," in row_txt or "A" in row_txt:
            for row in table.rows:
                rt = row.cells[0].text
                if "................" in rt and "A," not in rt:
                    if mcq_idx < len(mcq_data):
                        q = mcq_data[mcq_idx]['q']
                        for cell in row.cells:
                            for p in cell.paragraphs:
                                if "................" in p.text:
                                    p.text = p.text.replace("....................................................", q).replace("..................", "")
                elif ("A," in rt or "A" in rt) and mcq_idx < len(mcq_data):
                    opts = mcq_data[mcq_idx]['opts']
                    random.shuffle(opts)
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            t = p.text
                            t = t.replace("A,..................", f"A, {opts[0]}")
                            t = t.replace("B,..................", f"B, {opts[1]}")
                            t = t.replace("C,..................", f"C, {opts[2]}")
                            t = t.replace("D,..................", f"D, {opts[3]}")
                            t = t.replace("E,..................", f"E, {opts[4]}")
                            p.text = t
                    mcq_idx += 1
        
        # TF Logic
        else:
            is_tf = False
            for r in table.rows:
                if "(" in r.cells[-1].text and ")" in r.cells[-1].text:
                    is_tf = True
                    break
            if is_tf:
                for r in table.rows:
                    if tf_idx < len(tf_data):
                        cell = r.cells[0]
                        replaced = False
                        for p in cell.paragraphs:
                            if "................" in p.text:
                                p.text = p.text.replace("....................................................", tf_data[tf_idx]).replace("..................", "")
                                replaced = True
                        if replaced: tf_idx += 1
    
    # Save to memory buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- واجهة التطبيق ---
uploaded_file = st.file_uploader("ارفع ملف بنك الأسئلة (Word)", type=['docx'])

if uploaded_file is not None:
    st.success("تم رفع الملف بنجاح!")
    
    if st.button("توليد الامتحان"):
        with st.spinner('جاري إعداد الأسئلة...'):
            try:
                mcq, tf = read_questions(uploaded_file)
                
                if not mcq and not tf:
                    st.error("لم يتم العثور على أسئلة! تأكد من التنسيق (# اختياري / # صح وخطأ).")
                else:
                    st.info(f"تم العثور على {len(mcq)} سؤال اختياري و {len(tf)} سؤال صح/خطأ.")
                    
                    # التوليد
                    final_file = generate_exam(mcq, tf, TEMPLATE_FILE)
                    
                    # زر التحميل
                    st.download_button(
                        label="📥 تحميل ورقة الامتحان الجاهزة",
                        data=final_file,
                        file_name="Exam_Paper.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
            except Exception as e:
                st.error(f"حدث خطأ: {e}")
                st.write("تأكد أن ملف القالب موجود في نفس المجلد باسم: " + TEMPLATE_FILE)
